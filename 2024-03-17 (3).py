from docx import Document
import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import pandas as pd
from openpyxl import load_workbook

load_wb = load_workbook('xlsx/1.xlsx')
load_ws = load_wb.active

name_list = []
birthday_list = []
ho_list = []

for i in range(1, load_ws.max_row + 1): #load_ws에 마지막 줄까지 반복

    name_list.append(load_ws.cell(i, 1).value)
    birthday_list.append(load_ws.cell(i, 2).value)
    ho_list.append(load_ws.cell(i, 3).value)

print(name_list)
print(birthday_list)
print(ho_list)

for i in range(len(name_list)):
    doc = docx.Document('docx/수료증양식.docx')
    style = doc.styles['Normal']
    style.font.name = '나눔고딕'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(12)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('제 ' + ho_list[i] + '호\n')
    style.font.name = '나눔고딕'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('수 료 증')
    style.font.name = '나눔고딕'
    run.blod = True
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(40)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run(' 성 명: ' + name_list[i] +'\n')
    style.font.name = '나눔고딕'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(20)
    run = para.add_run(' 생 년 월 일: ' + birthday_list[i] +'\n')
    style.font.name = '나눔고딕'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(20)
    run = para.add_run(' 교 육 과 정: 파이썬과 40개의 작품들 \n')

    doc.save('docx/'+name_list[i]+'.docx')
    convert('docx/'+name_list[i]+'.docx', 
            'docx/'+name_list[i]+'.pdf')



df = pd.DataFrame([["엄준식", "2999.03.29", "2024-0003"],
                   ["타다톤", "2929.04.19", "2024-0004"],
                   ["김춘식", "2999.07.39", "2024-0005"],
                   ["기모찌", "2949.11.02", "2024-0006"],
                   ["아아아", "2949.12.03", "2024-0007"]])

print(df)
df.to_excel('xlsx/2.xlsx', index=False, header=False)

